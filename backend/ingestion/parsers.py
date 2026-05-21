"""
ingestion/parsers.py
--------------------
PURPOSE:
    Extracts plain text from different document formats.
    Every parser receives a source (file path or URL) and returns a single
    clean string of text. The rest of the pipeline (chunker, embedder) never
    needs to know what format the original document was in.

CONCEPT — Why separate parsing from chunking?
    Separation of concerns: parsers deal with format complexity (binary PDF
    structure, DOCX XML, HTML tags). Once we have plain text, the chunking
    logic is identical regardless of source format.

SUPPORTED FORMATS:
    PDF   — pdfplumber extracts text page by page, handling columns and tables
            better than PyPDF2 for most documents.
    DOCX  — python-docx reads paragraph objects; we join them with newlines.
    TXT / MD — plain read; no library needed.
    URL   — requests fetches HTML; BeautifulSoup strips tags and extracts body text.

ADDING A NEW FORMAT:
    Write a function matching the signature `parse_xxx(source: str) -> str`
    and add it to the `parse_document` dispatcher at the bottom.

USAGE:
    from ingestion.parsers import parse_document
    text = parse_document("/path/to/policy.pdf")
    text = parse_document("https://internal-wiki.company.com/hr-policy")
"""

import re
from pathlib import Path
from core.logging_config import get_logger

logger = get_logger(__name__)


# ── PDF ───────────────────────────────────────────────────────────────────────

def parse_pdf(file_path: str) -> str:
    """
    Extract text from a PDF file page by page using pdfplumber.

    pdfplumber is preferred over PyPDF2 because it handles multi-column
    layouts, embedded tables, and rotated text more reliably.

    Args:
        file_path: Absolute or relative path to the .pdf file.

    Returns:
        Concatenated text from all pages, separated by newlines.
        Empty pages are skipped silently.
    """
    import pdfplumber  # imported here so missing library gives a clear error

    logger.info("Parsing PDF: %s", file_path)
    pages_text = []

    with pdfplumber.open(file_path) as pdf:
        total_pages = len(pdf.pages)
        logger.debug("PDF has %d pages", total_pages)

        for i, page in enumerate(pdf.pages):
            text = page.extract_text()
            if text and text.strip():
                pages_text.append(text.strip())
            else:
                logger.debug("Page %d/%d is empty or image-only — skipped", i + 1, total_pages)

    full_text = "\n\n".join(pages_text)
    logger.info("PDF parsed: %d pages extracted, %d characters total", len(pages_text), len(full_text))
    return full_text


# ── DOCX ──────────────────────────────────────────────────────────────────────

def parse_docx(file_path: str) -> str:
    """
    Extract text from a Word (.docx) document paragraph by paragraph.

    python-docx reads the Open XML structure. Each paragraph becomes one line.
    Tables are not extracted in this implementation — add table iteration
    from doc.tables if table content is important for your policies.

    Args:
        file_path: Absolute or relative path to the .docx file.

    Returns:
        Text of all non-empty paragraphs joined by newlines.
    """
    from docx import Document

    logger.info("Parsing DOCX: %s", file_path)
    doc = Document(file_path)

    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    full_text = "\n".join(paragraphs)

    logger.info("DOCX parsed: %d paragraphs, %d characters total", len(paragraphs), len(full_text))
    return full_text


# ── TXT / Markdown ────────────────────────────────────────────────────────────

def parse_text(file_path: str) -> str:
    """
    Read a plain text or Markdown file directly.

    UTF-8 is assumed. If the file uses a different encoding, pass the
    encoding explicitly or add charset detection (chardet library).

    Args:
        file_path: Absolute or relative path to the .txt or .md file.

    Returns:
        Raw file content as a string.
    """
    logger.info("Parsing text file: %s", file_path)
    text = Path(file_path).read_text(encoding="utf-8")
    logger.info("Text file parsed: %d characters", len(text))
    return text


# ── URL / HTML ────────────────────────────────────────────────────────────────

def parse_url(url: str) -> str:
    """
    Fetch a web page and extract its visible text content.

    Strategy:
      1. requests.get() fetches the raw HTML.
      2. BeautifulSoup parses it and removes <script> and <style> tags,
         which contain code/CSS, not readable content.
      3. get_text() extracts visible text with whitespace collapsed.

    Args:
        url: Full HTTP/HTTPS URL.

    Returns:
        Cleaned visible text from the page body.

    Raises:
        requests.RequestException if the page cannot be fetched.
    """
    import requests
    from bs4 import BeautifulSoup

    logger.info("Fetching URL: %s", url)
    headers = {"User-Agent": "Mozilla/5.0 (enterprise-chatbot-indexer/1.0)"}
    response = requests.get(url, headers=headers, timeout=15)
    response.raise_for_status()  # raises HTTPError for 4xx/5xx status codes

    soup = BeautifulSoup(response.text, "html.parser")

    # Remove non-content tags before extracting text.
    for tag in soup(["script", "style", "nav", "footer", "header"]):
        tag.decompose()

    text = soup.get_text(separator="\n", strip=True)

    # Collapse runs of blank lines left behind after tag removal.
    text = re.sub(r"\n{3,}", "\n\n", text)

    logger.info("URL parsed: %d characters extracted from %s", len(text), url)
    return text


# ── Dispatcher ────────────────────────────────────────────────────────────────

def parse_document(source: str) -> str:
    """
    Route `source` to the correct parser based on its type.

    Args:
        source: File path (PDF/DOCX/TXT/MD) or HTTP/HTTPS URL.

    Returns:
        Extracted plain text.

    Raises:
        ValueError if the file extension is not supported.
        FileNotFoundError if the file path does not exist.
    """
    if source.startswith("http://") or source.startswith("https://"):
        return parse_url(source)

    path = Path(source)
    if not path.exists():
        raise FileNotFoundError(f"Document not found: {source}")

    ext = path.suffix.lower()
    parsers = {
        ".pdf": parse_pdf,
        ".docx": parse_docx,
        ".doc": parse_docx,    # older Word format — python-docx handles most .doc files
        ".txt": parse_text,
        ".md": parse_text,
    }

    if ext not in parsers:
        raise ValueError(
            f"Unsupported file type: '{ext}'. "
            f"Supported: {', '.join(parsers.keys())} and HTTP/HTTPS URLs."
        )

    return parsers[ext](source)
